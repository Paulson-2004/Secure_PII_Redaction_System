from io import BytesIO

from docx import Document

from main import _read_docx_text_bytes


def test_docx_in_memory_extraction():
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("Second line")

    buf = BytesIO()
    doc.save(buf)

    text = _read_docx_text_bytes(buf.getvalue())
    assert "Hello World" in text
    assert "Second line" in text
