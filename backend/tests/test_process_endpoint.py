from io import BytesIO
import os

from fastapi.testclient import TestClient
from docx import Document
from encryption import generate_key


def test_process_endpoint_txt(app_factory):
    client = TestClient(app_factory())
    data = b"Name: John Doe\nEmail: john@gmail.com\nPhone: 9876543210"
    files = {"file": ("sample.txt", BytesIO(data), "text/plain")}

    response = client.post("/process/", files=files)
    assert response.status_code == 200

    payload = response.json()
    assert payload["total_pii_detected"] >= 1
    assert "redacted_text" in payload


def test_process_endpoint_return_pdf(app_factory):
    client = TestClient(app_factory())
    data = b"Email: john@gmail.com"
    files = {"file": ("sample.txt", BytesIO(data), "text/plain")}

    response = client.post("/process/?return_pdf=true", files=files)
    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")


def test_process_endpoint_docx(app_factory):
    client = TestClient(app_factory())

    doc = Document()
    doc.add_paragraph("Name: John Doe")
    doc.add_paragraph("Email: john@gmail.com")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    files = {
        "file": (
            "sample.docx",
            buf,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    response = client.post("/process/", files=files)
    assert response.status_code == 200
    payload = response.json()
    assert payload["total_pii_detected"] >= 1


def test_process_endpoint_rejects_content_type(app_factory):
    client = TestClient(app_factory())
    data = b"test"
    files = {"file": ("sample.txt", BytesIO(data), "application/x-msdownload")}

    response = client.post("/process/", files=files)
    assert response.status_code == 400


def test_process_endpoint_rejects_extension(app_factory):
    client = TestClient(app_factory())
    data = b"test"
    files = {"file": ("sample.exe", BytesIO(data), "application/octet-stream")}

    response = client.post("/process/", files=files)
    assert response.status_code == 400


def test_process_endpoint_rejects_large_upload(app_factory):
    client = TestClient(app_factory({"APP_MAX_UPLOAD_MB": "0"}))
    data = b"large"
    files = {"file": ("sample.txt", BytesIO(data), "text/plain")}

    response = client.post("/process/", files=files)
    assert response.status_code == 413


def test_decrypt_endpoint_success(app_factory, tmp_path):
    key = generate_key()
    app = app_factory(
        {
            "APP_ENABLE_DECRYPT_ENDPOINT": "true",
            "APP_ADMIN_TOKEN": "secret",
            "APP_ENCRYPTION_ENABLED": "true",
            "APP_ENCRYPTION_KEY": key,
            "APP_UPLOADS_DIR": str(tmp_path / "uploads"),
        }
    )
    client = TestClient(app)

    from encryption import encrypt_bytes

    uploads_dir = os.getenv("APP_UPLOADS_DIR")
    os.makedirs(uploads_dir, exist_ok=True)
    filename = "sample.txt"
    original = b"hello world"
    encrypted = encrypt_bytes(original)
    with open(os.path.join(uploads_dir, filename), "wb") as f:
        f.write(encrypted)

    response = client.get(f"/decrypt?filename={filename}&token=secret")
    assert response.status_code == 200
    assert response.content == original


def test_decrypt_endpoint_rejects_bad_token(app_factory, tmp_path):
    key = generate_key()
    app = app_factory(
        {
            "APP_ENABLE_DECRYPT_ENDPOINT": "true",
            "APP_ADMIN_TOKEN": "secret",
            "APP_ENCRYPTION_ENABLED": "true",
            "APP_ENCRYPTION_KEY": key,
            "APP_UPLOADS_DIR": str(tmp_path / "uploads"),
        }
    )
    client = TestClient(app)
    response = client.get("/decrypt?filename=sample.txt&token=wrong")
    assert response.status_code == 403


def test_decrypt_endpoint_requires_encryption(app_factory, tmp_path):
    app = app_factory(
        {
            "APP_ENABLE_DECRYPT_ENDPOINT": "true",
            "APP_ADMIN_TOKEN": "secret",
            "APP_ENCRYPTION_ENABLED": "false",
            "APP_UPLOADS_DIR": str(tmp_path / "uploads"),
        }
    )
    client = TestClient(app)
    response = client.get("/decrypt?filename=sample.txt&token=secret")
    assert response.status_code == 400
