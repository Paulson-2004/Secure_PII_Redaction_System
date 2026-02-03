import io
import os
import re
import smtplib
import uuid
from typing import List, Optional
from email.message import EmailMessage

from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel

from config import CONFIG
from datetime import datetime

from db import (
    RedactionLogData,
    create_user,
    create_password_reset_token,
    change_password,
    fetch_log_by_id,
    fetch_logs,
    get_user_by_token,
    log_redaction,
    login_user,
    logout_user,
    reset_password_admin,
    reset_password_with_token,
)
from ocr import extract_text_and_boxes
from pdf_generator import generate_redacted_pdf
from pii_detector import detect_pii
from redaction import redact_text
from encryption import decrypt_bytes, encrypt_bytes
from media_redaction import redact_image_bytes, redact_pdf_with_boxes
from docx import Document

app = FastAPI()

os.makedirs(CONFIG.uploads_dir, exist_ok=True)
os.makedirs(CONFIG.output_dir, exist_ok=True)


class UserCredentials(BaseModel):
    username: str
    password: str


class RegisterRequest(BaseModel):
    username: str
    password: str
    email: Optional[str] = None


class ChangePasswordRequest(BaseModel):
    old_password: str
    new_password: str


class ResetPasswordRequest(BaseModel):
    username: str
    new_password: str


class ForgotPasswordRequest(BaseModel):
    email: str


class ResetPasswordTokenRequest(BaseModel):
    token: str
    new_password: str



def _safe_filename(original_name: str) -> str:
    ext = os.path.splitext(original_name)[1].lower()
    if ext not in set(CONFIG.allowed_extensions):
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")
    return f"{uuid.uuid4().hex}{ext}"


def _validate_content_type(content_type: Optional[str]) -> None:
    if not content_type:
        raise HTTPException(status_code=400, detail="Missing content type")
    base_type = content_type.split(";", 1)[0].strip().lower()
    allowed = {item.strip().lower() for item in CONFIG.allowed_content_types}
    if base_type not in allowed:
        raise HTTPException(status_code=400, detail=f"Unsupported content type: {content_type}")


def _enforce_size_limit(data: bytes) -> None:
    max_bytes = CONFIG.max_upload_mb * 1024 * 1024
    if len(data) > max_bytes:
        raise HTTPException(status_code=413, detail="File too large")


def _normalize_token(token: str) -> str:
    return token.strip().strip(".,;:()[]{}<>\"'").lower()


def _find_matching_indices(words: List[dict], pii_value: str) -> List[int]:
    tokens = [t for t in (_normalize_token(t) for t in pii_value.split()) if t]
    word_norms = [_normalize_token(w["text"]) for w in words]

    if not tokens:
        return []

    # Exact sequence match
    if len(tokens) > 1:
        for i in range(0, len(word_norms) - len(tokens) + 1):
            if word_norms[i : i + len(tokens)] == tokens:
                return list(range(i, i + len(tokens)))

    # Single-token match (or fallback)
    indices = [i for i, w in enumerate(word_norms) if w == tokens[0]]
    if indices:
        return indices

    # Fallback: substring match
    for i, w in enumerate(words):
        if tokens[0] in _normalize_token(w["text"]):
            return [i]

    return []


def _validate_password(password: str) -> Optional[str]:
    if len(password) < 8:
        return "Password must be at least 8 characters"
    if not re.search(r"[A-Z]", password):
        return "Password must include an uppercase letter"
    if not re.search(r"[a-z]", password):
        return "Password must include a lowercase letter"
    if not re.search(r"[0-9]", password):
        return "Password must include a number"
    if not re.search(r"[^A-Za-z0-9]", password):
        return "Password must include a special character"
    return None


def _normalize_email(email: Optional[str]) -> Optional[str]:
    if not email:
        return None
    cleaned = email.strip().lower()
    return cleaned if cleaned else None


def _validate_email(email: str) -> bool:
    return bool(re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email))


def _send_reset_email(to_email: str, token: str) -> None:
    host = CONFIG.smtp_host
    port = CONFIG.smtp_port
    from_addr = CONFIG.smtp_from or CONFIG.smtp_user
    if not host or not from_addr:
        raise HTTPException(status_code=500, detail="SMTP not configured")

    msg = EmailMessage()
    msg["Subject"] = "Secure PII Password Reset"
    msg["From"] = from_addr
    msg["To"] = to_email
    msg.set_content(
        "\n".join(
            [
                "You requested a password reset.",
                f"Your reset token is: {token}",
                f"This token expires in {CONFIG.reset_token_ttl_minutes} minutes.",
                "If you did not request this, you can ignore this email.",
            ]
        )
    )

    try:
        if CONFIG.smtp_use_tls:
            with smtplib.SMTP(host, port) as server:
                server.starttls()
                if CONFIG.smtp_user and CONFIG.smtp_password:
                    server.login(CONFIG.smtp_user, CONFIG.smtp_password)
                server.send_message(msg)
        else:
            with smtplib.SMTP(host, port) as server:
                if CONFIG.smtp_user and CONFIG.smtp_password:
                    server.login(CONFIG.smtp_user, CONFIG.smtp_password)
                server.send_message(msg)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"SMTP send failed: {exc}")


@app.get("/health")
def health_check():
    return {"status": "ok"}


@app.post("/auth/register")
def register_user(creds: RegisterRequest):
    password_error = _validate_password(creds.password)
    if password_error:
        raise HTTPException(status_code=400, detail=password_error)
    email = _normalize_email(creds.email)
    if email and not _validate_email(email):
        raise HTTPException(status_code=400, detail="Invalid email format")
    try:
        user = create_user(creds.username.strip(), creds.password, email=email)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    if not user:
        raise HTTPException(status_code=500, detail="Database not configured")
    return user


@app.post("/auth/login")
def login(creds: UserCredentials):
    user = login_user(creds.username.strip(), creds.password)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid username or password")
    return user


@app.get("/auth/me")
def whoami(user_token: Optional[str] = None):
    user = _require_user(user_token)
    return user


@app.post("/auth/logout")
def logout(user_token: Optional[str] = None):
    _require_user(user_token)
    ok = logout_user(user_token or "")
    if not ok:
        raise HTTPException(status_code=400, detail="Logout failed")
    return {"status": "ok"}


@app.post("/auth/change-password")
def change_password_endpoint(
    payload: ChangePasswordRequest, user_token: Optional[str] = None
):
    _require_user(user_token)
    password_error = _validate_password(payload.new_password)
    if password_error:
        raise HTTPException(status_code=400, detail=password_error)
    ok = change_password(user_token or "", payload.old_password, payload.new_password)
    if not ok:
        raise HTTPException(status_code=400, detail="Invalid password")
    return {"status": "ok"}


@app.post("/auth/reset-password")
def reset_password(payload: ResetPasswordRequest, token: Optional[str] = None):
    _require_admin_token(token)
    password_error = _validate_password(payload.new_password)
    if password_error:
        raise HTTPException(status_code=400, detail=password_error)
    ok = reset_password_admin(payload.username.strip(), payload.new_password)
    if not ok:
        raise HTTPException(status_code=400, detail="Invalid username")
    return {"status": "ok"}


@app.post("/auth/forgot-password")
def forgot_password(payload: ForgotPasswordRequest):
    email = _normalize_email(payload.email)
    if not email:
        raise HTTPException(status_code=400, detail="Email is required")
    if not _validate_email(email):
        raise HTTPException(status_code=400, detail="Invalid email format")
    if not CONFIG.db_url:
        raise HTTPException(status_code=500, detail="Database not configured")
    if not CONFIG.smtp_host or not (CONFIG.smtp_from or CONFIG.smtp_user):
        raise HTTPException(status_code=500, detail="SMTP not configured")
    token = create_password_reset_token(email)
    if token:
        _send_reset_email(email, token)
    return {"status": "ok"}


@app.post("/auth/reset-password-token")
def reset_password_token(payload: ResetPasswordTokenRequest):
    password_error = _validate_password(payload.new_password)
    if password_error:
        raise HTTPException(status_code=400, detail=password_error)
    if not CONFIG.db_url:
        raise HTTPException(status_code=500, detail="Database not configured")
    ok = reset_password_with_token(payload.token.strip(), payload.new_password)
    if not ok:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    return {"status": "ok"}


@app.get("/config")
def config_debug(token: Optional[str] = None):
    _require_api_token(token)
    if not CONFIG.enable_config_debug:
        raise HTTPException(status_code=404, detail="Not found")

    return {
        "allowed_extensions": CONFIG.allowed_extensions,
        "allowed_content_types": CONFIG.allowed_content_types,
        "uploads_dir": CONFIG.uploads_dir,
        "output_dir": CONFIG.output_dir,
        "enable_config_debug": CONFIG.enable_config_debug,
        "max_upload_mb": CONFIG.max_upload_mb,
        "enable_rag_stub": CONFIG.enable_rag_stub,
        "rag_vectordb_url": CONFIG.rag_vectordb_url,
        "encryption_enabled": CONFIG.encryption_enabled,
        "enable_decrypt_endpoint": CONFIG.enable_decrypt_endpoint,
        "use_preprocess": CONFIG.use_preprocess,
        "pdf_dpi": CONFIG.pdf_dpi,
        "ner_model_path": CONFIG.ner_model_path,
        "db_url": CONFIG.db_url,
        "db_echo": CONFIG.db_echo,
    }


def _parse_date(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid date format: {value}")


@app.get("/logs")
def get_logs(
    limit: int = 100,
    offset: int = 0,
    filename: Optional[str] = None,
    pii_type: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    sort_by: Optional[str] = None,
    sort_dir: Optional[str] = None,
    token: Optional[str] = None,
    user_token: Optional[str] = None,
):
    _require_api_token(token)
    if CONFIG.admin_token:
        _require_admin_token(token)
    user = _resolve_user(user_token)
    limit = max(1, min(limit, 1000))
    offset = max(0, offset)
    rows, total = fetch_logs(
        limit=limit,
        offset=offset,
        filename=filename,
        pii_type=pii_type,
        date_from=_parse_date(date_from),
        date_to=_parse_date(date_to),
        sort_by=sort_by,
        sort_dir=sort_dir,
        user_id=user["id"] if user else None,
    )
    return {
        "count": len(rows),
        "count_total": total,
        "limit": limit,
        "offset": offset,
        "logs": rows,
    }


@app.get("/logs/{log_id}")
def get_log_by_id(log_id: int, token: Optional[str] = None, user_token: Optional[str] = None):
    _require_api_token(token)
    if CONFIG.admin_token:
        _require_admin_token(token)
    row = fetch_log_by_id(log_id)
    if not row:
        raise HTTPException(status_code=404, detail="Log not found")
    user = _resolve_user(user_token)
    if user and row.get("user_id") and row.get("user_id") != user["id"]:
        raise HTTPException(status_code=404, detail="Log not found")
    return row


def _read_docx_text(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paragraphs)


def _read_docx_text_bytes(data: bytes) -> str:
    with io.BytesIO(data) as buf:
        doc = Document(buf)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paragraphs)


def _require_admin_token(token: Optional[str]) -> None:
    if not CONFIG.admin_token:
        raise HTTPException(status_code=403, detail="Admin token not configured")
    if not token or token != CONFIG.admin_token:
        raise HTTPException(status_code=403, detail="Invalid admin token")


def _require_api_token(token: Optional[str]) -> None:
    if not CONFIG.api_token:
        return
    if not token or token != CONFIG.api_token:
        raise HTTPException(status_code=401, detail="Invalid API token")


def _resolve_user(user_token: Optional[str]) -> Optional[dict]:
    if not user_token:
        return None
    user = get_user_by_token(user_token)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid user token")
    return user


def _require_user(user_token: Optional[str]) -> dict:
    if not user_token:
        raise HTTPException(status_code=401, detail="user_token required")
    return _resolve_user(user_token)


@app.get("/decrypt")
def decrypt_file(filename: str, token: Optional[str] = None):
    _require_api_token(token)
    if not CONFIG.enable_decrypt_endpoint:
        raise HTTPException(status_code=404, detail="Not found")
    _require_admin_token(token)
    if not CONFIG.encryption_enabled:
        raise HTTPException(status_code=400, detail="Encryption not enabled")

    safe_name = os.path.basename(filename)
    path = os.path.join(CONFIG.uploads_dir, safe_name)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="File not found")

    with open(path, "rb") as f:
        encrypted = f.read()

    try:
        decrypted = decrypt_bytes(encrypted)
    except ValueError as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    return Response(
        content=decrypted,
        media_type="application/octet-stream",
        headers={"Content-Disposition": f"attachment; filename=decrypted_{safe_name}"},
    )


@app.post("/process/")
async def process_file(
    file: UploadFile = File(...),
    return_pdf: bool = False,
    return_redacted_file: bool = False,
    token: Optional[str] = None,
    user_token: Optional[str] = None,
):
    _require_api_token(token)
    user = _resolve_user(user_token)
    _validate_content_type(file.content_type)
    safe_name = _safe_filename(file.filename)
    path = os.path.join(CONFIG.uploads_dir, safe_name)

    data = await file.read()
    _enforce_size_limit(data)

    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        text = data.decode("utf-8", errors="ignore")
        words = []
    elif ext == ".docx":
        text = _read_docx_text_bytes(data)
        words = []
    else:
        temp_path = path + ".tmp"
        with open(temp_path, "wb") as buffer:
            buffer.write(data)

        text, words = extract_text_and_boxes(temp_path, use_preprocess=CONFIG.use_preprocess)

        if os.path.exists(temp_path):
            os.remove(temp_path)

    if CONFIG.encryption_enabled:
        try:
            encrypted = encrypt_bytes(data)
        except ValueError as exc:
            raise HTTPException(status_code=500, detail=str(exc))
        with open(path, "wb") as buffer:
            buffer.write(encrypted)
    else:
        with open(path, "wb") as buffer:
            buffer.write(data)

    pii_data = detect_pii(text)
    redacted_text = redact_text(text, pii_data)

    boxes = []
    for pii in pii_data:
        indices = _find_matching_indices(words, pii["value"])
        for idx in indices:
            word = words[idx]
            boxes.append(
                {
                    "x": word["x"],
                    "y": word["y"],
                    "w": word["w"],
                    "h": word["h"],
                    "page": word.get("page", 0),
                    "type": pii["type"],
                    "start": word.get("start"),
                    "end": word.get("end"),
                }
            )

    pii_counts = {}
    for item in pii_data:
        pii_counts[item["type"]] = pii_counts.get(item["type"], 0) + 1
    try:
        log_redaction(
            RedactionLogData(
                user_id=user["id"] if user else None,
                username=user["username"] if user else None,
                filename=safe_name,
                content_type=file.content_type or "",
                size_bytes=len(data),
                total_pii=len(pii_data),
                pii_counts=pii_counts,
            )
        )
    except Exception:
        pass

    if return_pdf:
        output_path = os.path.join(CONFIG.output_dir, f"{uuid.uuid4().hex}.pdf")
        generate_redacted_pdf(redacted_text, output_path=output_path)
        return FileResponse(
            output_path,
            media_type="application/pdf",
            filename="redacted.pdf",
        )

    if return_redacted_file:
        if ext in {".png", ".jpg", ".jpeg"}:
            redacted_bytes = redact_image_bytes(data, boxes)
            return Response(
                content=redacted_bytes,
                media_type="image/png",
                headers={"Content-Disposition": "attachment; filename=redacted.png"},
            )
        if ext == ".pdf":
            temp_pdf = path + ".tmp.pdf"
            with open(temp_pdf, "wb") as f:
                f.write(data)
            output_path = os.path.join(CONFIG.output_dir, f"{uuid.uuid4().hex}.pdf")
            redact_pdf_with_boxes(temp_pdf, boxes, output_path=output_path)
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)
            return FileResponse(
                output_path,
                media_type="application/pdf",
                filename="redacted.pdf",
            )
        raise HTTPException(status_code=400, detail="Redaction file output not supported for this type")

    return {
        "total_pii_detected": len(pii_data),
        "redacted_text": redacted_text,
        "boxes": boxes,
        "pii": pii_data,
    }
